#!/usr/bin/env python3
"""Convert our markdown deliverables to nicely-formatted .docx files."""

import re
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def parse_inline(text):
    """Return list of (text, bold, italic) tuples from a single line."""
    parts = []
    # Token regex: **bold**, *italic*, `code`, [link](url), plain
    pattern = re.compile(
        r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^\)]+\))'
    )
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            parts.append((text[pos:m.start()], False, False))
        tok = m.group(0)
        if tok.startswith('**'):
            parts.append((tok[2:-2], True, False))
        elif tok.startswith('*'):
            parts.append((tok[1:-1], False, True))
        elif tok.startswith('`'):
            parts.append((tok[1:-1], False, True))
        elif tok.startswith('['):
            inner = re.match(r'\[([^\]]+)\]\(([^\)]+)\)', tok)
            if inner:
                parts.append((inner.group(1), False, False))
        pos = m.end()
    if pos < len(text):
        parts.append((text[pos:], False, False))
    return parts if parts else [(text, False, False)]


def add_formatted_paragraph(doc, text, style=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    for chunk, bold, italic in parse_inline(text):
        run = p.add_run(chunk)
        run.bold = bold
        run.italic = italic
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
    return p


def convert(md_path, docx_path, title=None, subtitle=None):
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Strip YAML frontmatter
    if content.startswith('---'):
        end = content.find('---', 3)
        if end != -1:
            content = content[end + 3:].lstrip()

    lines = content.split('\n')
    in_table = False
    table_rows = []
    in_code = False
    code_lines = []
    in_list = False

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip()

        # Fenced code block
        if stripped.startswith('```'):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # Table
        if '|' in stripped and stripped.startswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            # skip separator rows
            if re.match(r'^\|[\s\-:|]+\|\s*$', stripped):
                i += 1
                continue
            cells = [c.strip() for c in stripped.strip('|').split('|')]
            table_rows.append(cells)
            i += 1
            continue
        elif in_table:
            # flush
            in_table = False
            if table_rows:
                ncols = max(len(r) for r in table_rows)
                tbl = doc.add_table(rows=len(table_rows), cols=ncols)
                tbl.style = 'Light Grid Accent 1'
                for ri, row in enumerate(table_rows):
                    for ci in range(ncols):
                        cell = tbl.rows[ri].cells[ci]
                        text = row[ci] if ci < len(row) else ''
                        cell.text = ''
                        para = cell.paragraphs[0]
                        for chunk, bold, italic in parse_inline(text):
                            run = para.add_run(chunk)
                            run.bold = bold or (ri == 0)
                            run.italic = italic
                            run.font.size = Pt(10)
                table_rows = []

        # Horizontal rule
        if stripped in ('---', '***', '___'):
            doc.add_paragraph().add_run('').add_break()
            i += 1
            continue

        # Headings
        m = re.match(r'^(#{1,6})\s+(.*)$', stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip().rstrip('#').strip()
            h = doc.add_heading(level=min(level, 4))
            for chunk, bold, italic in parse_inline(text):
                run = h.add_run(chunk)
                run.bold = True
                run.italic = italic
                if level == 1:
                    run.font.size = Pt(22)
                elif level == 2:
                    run.font.size = Pt(15)
                elif level == 3:
                    run.font.size = Pt(12.5)
                else:
                    run.font.size = Pt(11)
            i += 1
            continue

        # Blockquote
        if stripped.startswith('> '):
            text = stripped[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            for chunk, bold, italic in parse_inline(text):
                run = p.add_run(chunk)
                run.bold = bold
                run.italic = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # Unordered list
        if re.match(r'^\s*[-*]\s+', line):
            text = re.sub(r'^\s*[-*]\s+', '', line)
            p = doc.add_paragraph(style='List Bullet')
            for chunk, bold, italic in parse_inline(text):
                run = p.add_run(chunk)
                run.bold = bold
                run.italic = italic
                run.font.size = Pt(11)
            i += 1
            continue

        # Ordered list
        if re.match(r'^\s*\d+\.\s+', line):
            text = re.sub(r'^\s*\d+\.\s+', '', line)
            p = doc.add_paragraph(style='List Number')
            for chunk, bold, italic in parse_inline(text):
                run = p.add_run(chunk)
                run.bold = bold
                run.italic = italic
                run.font.size = Pt(11)
            i += 1
            continue

        # Blank line
        if not stripped:
            i += 1
            continue

        # Regular paragraph
        add_formatted_paragraph(doc, stripped)
        i += 1

    # Flush trailing table
    if in_table and table_rows:
        ncols = max(len(r) for r in table_rows)
        tbl = doc.add_table(rows=len(table_rows), cols=ncols)
        tbl.style = 'Light Grid Accent 1'
        for ri, row in enumerate(table_rows):
            for ci in range(ncols):
                cell = tbl.rows[ri].cells[ci]
                text = row[ci] if ci < len(row) else ''
                cell.text = ''
                para = cell.paragraphs[0]
                for chunk, bold, italic in parse_inline(text):
                    run = para.add_run(chunk)
                    run.bold = bold or (ri == 0)
                    run.italic = italic
                    run.font.size = Pt(10)

    doc.save(docx_path)
    print(f'Wrote {docx_path}')


if __name__ == '__main__':
    convert(
        '/home/ubuntu/reagan_homeschool_dashboard/deliverables/helpers-manual.md',
        '/home/ubuntu/reagan_homeschool_dashboard/deliverables/Reagan-Dashboard-Helpers-Manual.docx',
    )
    convert(
        '/home/ubuntu/reagan_homeschool_dashboard/deliverables/tutor-quick-guide.md',
        '/home/ubuntu/reagan_homeschool_dashboard/deliverables/Reagan-Dashboard-Tutor-Quick-Guide.docx',
    )
